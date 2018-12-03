import docx
import sys
import xlrd

def read_excel(excel_path='file.xls', sheet=0):
    try:
        data = xlrd.open_workbook(excel_path)
        return data.sheets()[sheet]
    except Exception as e:
        print(str(e))


def traversal_of_excel(excel_table_obj):
    excel_column_data = []
    nrows = excel_table_obj.nrows
    for i in range(1, nrows):
        excel_column_data.append(excel_table_obj.row_values(i))
    return excel_column_data


def read_docx(docx_path):
    doc = docx.Document(docx_path)
    return doc


def update_table_of_docx(excel_row, doc_obj, doc_path):
    update_list = [(0, 2), (1, 2), (4, 2), (4, 4), (5, 2), (5, 4), (8, 2), (8, 6), (9, 2)]
    for table in doc_obj.tables:  # 遍历所有表格
        for i in range(len(update_list)):
            r, c = update_list[i]
            updateInfo = str(excel_row[i+1]).replace('?','')
            #table.rows[r].cells[c].text = updateInfo
            run = table.rows[r].cells[c].paragraphs[0].add_run(updateInfo)
            run.font.name = '宋体'
            run.font.size = 120000
    if excel_row[1] == None:
        doc_output_path = doc_path.split('-')[0]+"-"+str(excel_row[0])+ '.docx'
    else:
        doc_output_path = doc_path.split('-')[0] +"-"+ str(excel_row[0]).replace('?','')+"-"+ str(excel_row[1]).replace('?','') + '.docx'
    doc_obj.save(doc_output_path)

def update_Paragraph_of_docx(excel_row, doc_obj):
    run = doc_obj.paragraphs[1].add_run(str(excel_row[1]).replace('?',''))
    run.font.name = '宋体'
    run.font.size = 155000

if len(sys.argv) < 3 or len(sys.argv) > 3:
    print(len(sys.argv))
    print(str(sys.argv))
    print("传入参数有问题，需要两条参数（Excel文件路径 与 Word文件路径）！")
    sys.exit()
else:
    excel_path, doc_path = sys.argv[1], sys.argv[2]
    excel_table_obj = read_excel(excel_path, 0)  # 读取Excel文件,返回Excel对象
    getExcelList = traversal_of_excel(excel_table_obj)
    for excelRow in getExcelList:
        try:
            docObj = read_docx(doc_path)  # 读取Word文件,返回Word对象
            update_Paragraph_of_docx(excelRow, docObj)
            update_table_of_docx(excelRow, docObj, doc_path)
        except Exception as e:
            print("程序处理有误：" + str(e))

# def main(excel_path,doc_path):
#     excel_table_obj = read_excel(excel_path, 0)  # 读取Excel文件,返回Excel对象
#     getExcelList = traversal_of_excel(excel_table_obj)
#     for excelRow in getExcelList:
#         try:
#             docObj = read_docx(doc_path)  # 读取Word文件,返回Word对象
#             update_Paragraph_of_docx(excelRow, docObj)
#             update_table_of_docx(excelRow, docObj, doc_path)
#         except Exception as e:
#             print("程序处理有误：" + str(e))
#
# if __name__ == '__main__':
#     excel_path = 'C:\\Users\\Administrator\\Desktop\\trans\\甲方信息表(1)(1).xlsx'
#     doc_path = 'C:\\Users\\Administrator\\Desktop\\trans\\通联支付网络服务股份有限公司当面付业务协议-模板(2).docx'
#     main(excel_path,doc_path)