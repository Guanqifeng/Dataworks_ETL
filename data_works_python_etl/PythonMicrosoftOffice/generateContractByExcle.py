import xlrd
import docx
import sys
def getExcel(exclepath,sheet = 0):
    getExcelRows = []
    try:
        data = xlrd.open_workbook(exclepath)
        table = data.sheets()[sheet]
        nrows = table.nrows
        for rownum in range(1, nrows):
            getExcelRows.append(table.row_values(rownum))
            #print(table.row_values(rownum))
        return getExcelRows
    except Exception as e:
        print("Excel处理异常："+str(e))
def connctWord(wordmodlepath):
    try:
        doc = docx.Document(wordmodlepath)
        return doc
    except Exception as e:
        print("Word获取异常：" + str(e))
def getContract(wordmodlepath,infolist):
    upadteIndex = [1,2,5]
    try:
        for row in infolist:
            getDoc = connctWord(wordmodlepath)
            for columnsindex in range(len(upadteIndex)):
                if upadteIndex[columnsindex] == 1:
                    p = getDoc.paragraphs[upadteIndex[columnsindex]].clear()
                    run1 = p.add_run("甲方："+str(row[0]))
                    font = run1.font
                    font.size = 155000
                    font.name = '宋体'
                elif upadteIndex[columnsindex] == 2:
                    p = getDoc.paragraphs[upadteIndex[columnsindex]].clear()
                    run1 = p.add_run("乙方："+str(row[1]))
                    font = run1.font
                    font.size = 155000
                    font.name = '宋体'
                elif upadteIndex[columnsindex] == 5:
                    p = getDoc.paragraphs[upadteIndex[columnsindex]].clear()
                    run1 = p.add_run('1、甲方将位于 {0} 商铺租给乙方使用，'
                              '交租方式按每季度交租，在租房期内，'
                              '乙方不得经营 违法活动 ，租用期限为    5   年'
                              '，（从 2015 年 1 月31日起至 2020 年 1 月  30 日止），'
                              '每月租金为 {1} 元，一季度租金为 {2} 元，'
                              '租金由乙方在交纳上一季度租金使用到期提前五天存入甲方账户（ {3} ，'
                              '账号：{4} ，户名：{5} ）。'
                              '先交租金后使用，逾期不交租金，甲方有权收回房屋，终止该房租赁合同，'
                              '没收乙方所交押金作为补偿甲方损失。'.format(row[2],row[3],row[4],row[5],row[6],row[7]))
                    #run1.underline = True
                    font = run1.font
                    font.size = 155000
                    font.name = '宋体'
            outPutPath = wordmodlepath.replace("模板.docx","")+"-"+str(row[0])+".docx"
            getDoc.save(outPutPath)
    except Exception as e:
        print("Word处理异常："+str(e))
def main(exclepath,wordmodlepath):
    getInfoList = getExcel(exclepath)
    getContract(wordmodlepath, getInfoList)

if len(sys.argv) < 3 or len(sys.argv) > 3:
    print(len(sys.argv))
    print(str(sys.argv))
    print("传入参数有问题，需要两条参数（Excel文件路径 与 Word文件路径）！")
    sys.exit()
else:
    excel_path, doc_path = sys.argv[1], sys.argv[2]
    main(excel_path, doc_path)
# if __name__ == '__main__':
#     filepath = 'C:\\Users\\Administrator\\Desktop\\test\\租赁合同信息.xlsx'
#     wordmodlepath = 'C:\\Users\\Administrator\\Desktop\\test\\租赁合同模板.docx'
#     main(filepath,wordmodlepath)